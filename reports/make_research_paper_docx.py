"""Generate the formal MBA-format research paper (Word .docx) for the
bank customer churn prediction project.

Formatting follows standard MBA project-report guidelines:
  - Times New Roman 12 pt body, 1.5 line spacing, justified
  - Headings 16/13/12 pt bold; chapter titles centered, new page each
  - A4 paper; margins: left 3.5 cm (binding), others 2.5 cm
  - Preliminary pages numbered in Roman numerals, chapters in Arabic
  - APA 7th edition references with hanging indent
  - Table of Contents / List of Figures / List of Tables as live Word
    fields: select all (Ctrl+A) and press F9 in Word to populate.

Run:  py -3.14 reports/make_research_paper_docx.py
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
FIGS = ROOT / "reports" / "figures"
OUT = ROOT / "reports" / "research_paper.docx"

BLACK = RGBColor(0, 0, 0)
FONT = "Times New Roman"


# ----------------------------------------------------------------- helpers
def set_style_font(style, size, bold=False, italic=False, name=FONT):
    style.font.name = name
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.italic = italic
    style.font.color.rgb = BLACK
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    for attr in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rfonts.set(qn(attr), name)


def add_field(paragraph, code, placeholder=""):
    """Insert a Word field (TOC, PAGE, SEQ...) with a placeholder result."""
    r1 = paragraph.add_run()._r
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
    r1.append(fld)
    r2 = paragraph.add_run()._r
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = code
    r2.append(instr)
    r3 = paragraph.add_run()._r
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "separate")
    r3.append(fld)
    paragraph.add_run(placeholder)
    r5 = paragraph.add_run()._r
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "end")
    r5.append(fld)


def add_footer_page_number(section, fmt=None, start=None):
    """Centered PAGE field in the footer; optional number format/restart."""
    sect_pr = section._sectPr
    for el in sect_pr.findall(qn("w:pgNumType")):
        sect_pr.remove(el)
    pg = OxmlElement("w:pgNumType")
    if fmt:
        pg.set(qn("w:fmt"), fmt)
    if start is not None:
        pg.set(qn("w:start"), str(start))
    sect_pr.append(pg)
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_field(p, " PAGE ", "1")
    for run in p.runs:
        run.font.name = FONT
        run.font.size = Pt(10)


def para(doc, text, style=None, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         bold=False, italic=False, size=None, space_after=6):
    p = doc.add_paragraph(style=style)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    run.font.name = FONT
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    return p


def prelim_title(doc, text):
    p = para(doc, text, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=14,
             space_after=18)
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(it)
        run.font.name = FONT
        run.font.size = Pt(12)
        p.paragraph_format.space_after = Pt(3)


def add_table(doc, headers, rows, caption=None, table_no=None,
              font_size=11):
    if caption:
        cap = doc.add_paragraph(style="Caption")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.add_run("Table ")
        add_field(cap, r" SEQ Table \* ARABIC ", str(table_no))
        cap.add_run(f": {caption}")
        for run in cap.runs:
            run.font.name = FONT
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.name = FONT
        run.font.size = Pt(font_size)
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "D9D9D9")
        cell._tc.get_or_add_tcPr().append(shd)
    for i, row in enumerate(rows, start=1):
        for j, val in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.name = FONT
            run.font.size = Pt(font_size)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return t


def add_figure(doc, filename, caption, fig_no, width=5.9):
    path = FIGS / filename
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = doc.add_paragraph(style="Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run("Figure ")
    add_field(cap, r" SEQ Figure \* ARABIC ", str(fig_no))
    cap.add_run(f": {caption}")
    for run in cap.runs:
        run.font.name = FONT


def chapter(doc, text):
    h = doc.add_heading(text, level=1)
    for run in h.runs:
        run.font.name = FONT
        run.font.color.rgb = BLACK
    return h


def section_h(doc, text):
    h = doc.add_heading(text, level=2)
    for run in h.runs:
        run.font.name = FONT
        run.font.color.rgb = BLACK
    return h


def reference(doc, text):
    p = doc.add_paragraph(style="APA Reference")
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(12)


# ----------------------------------------------------------------- document
doc = Document()

# base styles
normal = doc.styles["Normal"]
set_style_font(normal, 12)
normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
normal.paragraph_format.space_after = Pt(6)

h1 = doc.styles["Heading 1"]
set_style_font(h1, 16, bold=True)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.page_break_before = True
h1.paragraph_format.space_before = Pt(0)
h1.paragraph_format.space_after = Pt(18)
h1.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

h2 = doc.styles["Heading 2"]
set_style_font(h2, 13, bold=True)
h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h2.paragraph_format.space_before = Pt(12)
h2.paragraph_format.space_after = Pt(6)

h3 = doc.styles["Heading 3"]
set_style_font(h3, 12, bold=True, italic=True)

cap_style = doc.styles["Caption"]
set_style_font(cap_style, 10, italic=True)
cap_style.paragraph_format.space_before = Pt(4)
cap_style.paragraph_format.space_after = Pt(10)

lb = doc.styles["List Bullet"]
set_style_font(lb, 12)
lb.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

ref_style = doc.styles.add_style("APA Reference", 1)  # paragraph style
set_style_font(ref_style, 12)
ref_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
ref_style.paragraph_format.left_indent = Inches(0.5)
ref_style.paragraph_format.first_line_indent = Inches(-0.5)
ref_style.paragraph_format.space_after = Pt(8)
ref_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT

# page geometry - A4, MBA margins
sec1 = doc.sections[0]
sec1.page_width = Cm(21.0)
sec1.page_height = Cm(29.7)
sec1.left_margin = Cm(3.5)
sec1.right_margin = Cm(2.5)
sec1.top_margin = Cm(2.5)
sec1.bottom_margin = Cm(2.5)
sec1.different_first_page_header_footer = True  # no number on title page
add_footer_page_number(sec1, fmt="lowerRoman", start=1)

# ------------------------------------------------------------- title page
for _ in range(3):
    doc.add_paragraph()
para(doc, "PREDICTIVE MODELING AND RISK SCORING FOR",
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18, space_after=0)
para(doc, "BANK CUSTOMER CHURN",
     align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=18, space_after=14)
para(doc, "A Machine Learning Study of Customer Attrition in European Retail Banking",
     align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size=13, space_after=30)
para(doc, "A Research Paper", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
     size=12, space_after=30)
para(doc, "Submitted by", align=WD_ALIGN_PARAGRAPH.CENTER, size=12,
     space_after=2)
para(doc, "SHUBHANGI PANDEY", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True,
     size=14, space_after=60)
para(doc, "July 2026", align=WD_ALIGN_PARAGRAPH.CENTER, size=12)

# ------------------------------------------------------------- declaration
doc.add_page_break()
prelim_title(doc, "DECLARATION")
para(doc,
     "I hereby declare that this research paper, “Predictive Modeling and "
     "Risk Scoring for Bank Customer Churn,” is my own original work. "
     "All data analysis, model development, and interpretation "
     "presented herein were performed by me. Wherever the work of other authors "
     "has been used, it has been duly acknowledged and cited in accordance with "
     "APA (7th edition) referencing conventions. This paper has not been "
     "submitted, in whole or in part, for any other degree, diploma, or "
     "certification.")
for _ in range(3):
    doc.add_paragraph()
para(doc, "Shubhangi Pandey", align=WD_ALIGN_PARAGRAPH.RIGHT, bold=True,
     space_after=0)
para(doc, "July 2026", align=WD_ALIGN_PARAGRAPH.RIGHT)

# --------------------------------------------------------- acknowledgement
doc.add_page_break()
prelim_title(doc, "ACKNOWLEDGEMENT")
para(doc,
     "I am thankful to the open-source community behind Python, "
     "scikit-learn, XGBoost, SHAP, and Streamlit, whose tools made an "
     "end-to-end delivery — from raw data to a deployed decision-support "
     "dashboard — achievable within the project timeline. I also thank my "
     "peers and reviewers whose critical feedback materially improved the "
     "rigor and clarity of this paper.")

# ------------------------------------------------------- executive summary
doc.add_page_break()
prelim_title(doc, "EXECUTIVE SUMMARY")
para(doc,
     "Customer churn directly erodes the lifetime value of a retail-banking "
     "portfolio: in the 10,000-customer European dataset studied here, 20.37% "
     "of customers exited within the observation window. This research builds "
     "an end-to-end churn-prediction and risk-scoring system that allows the "
     "bank to move from reactive reporting to proactive, targeted retention.")
para(doc,
     "Five classification algorithms — Logistic Regression, Decision Tree, "
     "Random Forest, Gradient Boosting, and XGBoost — were benchmarked on a "
     "stratified 80/20 train–test split with five-fold cross-validated "
     "hyperparameter tuning. The tuned Gradient Boosting classifier was "
     "selected as the champion model, achieving a test ROC-AUC of 0.870, "
     "PR-AUC of 0.720, precision of 0.794, and F1-score of 0.602. Five "
     "engineered behavioral features (including an engagement score and "
     "balance-to-salary ratio) strengthened the signal available to all "
     "models.")
para(doc,
     "Explainability analysis using gain-based importance, SHAP values, and "
     "partial-dependence plots identified four dominant churn drivers: "
     "customer age (peaking in the 51–60 band at 56.2% churn), product "
     "holdings (customers with three or more products churn at 83–100%, "
     "while two-product customers churn at only 7.6%), account engagement "
     "(inactive members churn at nearly twice the rate of active members), "
     "and German residency (32.4% churn, double that of France or Spain).")
para(doc,
     "The paper closes with six managerial recommendations — including a "
     "two-product cross-sell guardrail, an activation campaign for inactive "
     "members, and a Germany-specific service diagnostic — and delivers a "
     "five-module Streamlit dashboard that operationalizes the model as an "
     "interactive risk calculator, portfolio scanner, and what-if simulator "
     "for retention teams.")

# ------------------------------------------------- TOC / lists (live fields)
doc.add_page_break()
prelim_title(doc, "TABLE OF CONTENTS")
p = doc.add_paragraph()
add_field(p, r' TOC \o "1-2" \h \z \u ',
          "Right-click here and choose “Update Field” to generate the "
          "Table of Contents.")
doc.add_page_break()
prelim_title(doc, "LIST OF FIGURES")
p = doc.add_paragraph()
add_field(p, r' TOC \h \z \c "Figure" ',
          "Right-click here and choose “Update Field” to generate the "
          "List of Figures.")
prelim_title(doc, "LIST OF TABLES")
p = doc.add_paragraph()
add_field(p, r' TOC \h \z \c "Table" ',
          "Right-click here and choose “Update Field” to generate the "
          "List of Tables.")

# ------------------------------------------- main section: arabic numbering
sec2 = doc.add_section(WD_SECTION.NEW_PAGE)
add_footer_page_number(sec2, fmt="decimal", start=1)

# ================================================== CHAPTER 1
chapter(doc, "CHAPTER 1: INTRODUCTION")

section_h(doc, "1.1 Background of the Study")
para(doc,
     "Retail banks operate in a mature, commoditized market in which the cost "
     "of acquiring a new customer substantially exceeds the cost of retaining "
     "an existing one; long-tenured customers additionally hold larger "
     "balances, adopt more products, and generate referral value (Reichheld & "
     "Sasser, 1990; Gupta, Lehmann, & Stuart, 2004). Customer churn — the "
     "voluntary closure or abandonment of a banking relationship — therefore "
     "attacks both current revenue and the compounding value of the customer "
     "base. The industry consensus that retention improvements of even a few "
     "percentage points translate into disproportionate profit gains has made "
     "churn management a board-level priority.")
para(doc,
     "Advances in machine learning have shifted churn management from "
     "retrospective analysis (“why did customers leave?”) to prospective "
     "risk scoring (“which current customers are likely to leave, and "
     "why?”). This study applies that prospective lens to a European "
     "retail-banking portfolio, combining predictive accuracy with model "
     "explainability so that scores can be trusted, audited, and acted upon "
     "by business stakeholders.")

section_h(doc, "1.2 Problem Statement")
para(doc,
     "The subject bank loses approximately one customer in five, yet its "
     "customer intelligence is descriptive rather than predictive: it cannot "
     "identify, ahead of time, which customers are at elevated risk, nor "
     "which attributes and behaviors drive that risk. Retention outreach is "
     "consequently untargeted, expensive, and late. The problem addressed "
     "here is the design, validation, and operationalization of a churn "
     "prediction model that produces accurate, calibrated, and explainable "
     "customer-level risk scores.")

section_h(doc, "1.3 Objectives of the Study")
para(doc, "Primary objectives:", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
bullets(doc, [
    "Predict customer churn with high discrimination (ROC-AUC) on held-out data.",
    "Produce a churn probability score for every customer to support risk banding.",
    "Identify and explain the dominant drivers of churn.",
])
para(doc, "Secondary objectives:", align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2)
bullets(doc, [
    "Reduce false-positive retention spend through a precision-oriented decision threshold.",
    "Make every prediction interpretable for compliance and business stakeholders.",
    "Enable scenario analysis through an interactive what-if simulation tool.",
])

section_h(doc, "1.4 Scope and Significance")
para(doc,
     "The study covers a single-snapshot portfolio of 10,000 customers across "
     "France, Germany, and Spain, with fourteen raw attributes spanning "
     "demographics, account status, and engagement. Its significance is "
     "twofold: managerially, it converts raw customer data into a deployed "
     "decision-support tool; methodologically, it demonstrates a disciplined "
     "benchmark-tune-explain workflow in which model selection is justified "
     "by out-of-sample evidence and every headline driver is corroborated by "
     "three independent explainability techniques.")

section_h(doc, "1.5 Organization of the Report")
para(doc,
     "Chapter 2 reviews the relevant literature. Chapter 3 details the "
     "research methodology. Chapter 4 presents the exploratory data "
     "analysis. Chapter 5 reports model development and comparative results. "
     "Chapter 6 presents model explainability. Chapter 7 consolidates "
     "findings and managerial recommendations. Chapter 8 discusses "
     "limitations and future scope, and Chapter 9 concludes.")

# ================================================== CHAPTER 2
chapter(doc, "CHAPTER 2: LITERATURE REVIEW")

section_h(doc, "2.1 The Economics of Customer Retention")
para(doc,
     "Reichheld and Sasser (1990) established the canonical argument that "
     "“zero defections” management materially improves profitability, "
     "estimating that a five-percentage-point improvement in retention can "
     "raise profits by 25–85% in service industries. Gupta, Lehmann, and "
     "Stuart (2004) formalized the link between retention and firm value, "
     "showing that a one-percent improvement in retention increases customer "
     "lifetime value substantially more than an equivalent improvement in "
     "acquisition cost or margin. These results motivate the managerial "
     "premium on early, accurate churn identification.")

section_h(doc, "2.2 Churn Prediction Methodologies")
para(doc,
     "Early churn models relied on logistic regression for its "
     "interpretability and calibrated probabilities. The comparative "
     "literature has since demonstrated consistent accuracy advantages for "
     "ensemble tree methods: Lemmens and Croux (2006) showed that bagging "
     "and boosting classification trees significantly outperform binary "
     "logit in churn prediction; Neslin et al. (2006), reporting on a "
     "large-scale churn tournament, found that tree-based ensembles "
     "dominated the leaderboard and that methodological choices materially "
     "affect the economic value of a churn model. Verbeke et al. (2012) "
     "extended the evaluation criterion from statistical accuracy to "
     "profit-driven measures, arguing that model selection should reflect "
     "the asymmetric costs of false positives (wasted retention spend) and "
     "false negatives (lost customers) — a consideration reflected in this "
     "study’s precision-oriented threshold policy.")
para(doc,
     "Among ensemble methods, Random Forests (Breiman, 2001) reduce variance "
     "through bootstrap aggregation of decorrelated trees, while gradient "
     "boosting (Friedman, 2001) builds additive stagewise models that "
     "typically achieve superior discrimination on structured tabular data. "
     "XGBoost (Chen & Guestrin, 2016) is a regularized, computationally "
     "efficient implementation of the same principle and a frequent "
     "benchmark winner on tabular problems.")

section_h(doc, "2.3 Explainable Machine Learning in Banking")
para(doc,
     "Predictive accuracy alone is insufficient in a regulated industry. "
     "SHAP (SHapley Additive exPlanations; Lundberg & Lee, 2017) provides "
     "theoretically grounded, locally accurate attributions of each "
     "prediction to each feature, enabling model risk management, adverse "
     "action reasoning, and stakeholder trust. Partial-dependence analysis "
     "complements SHAP by exposing the average shape of each "
     "feature–response relationship. This study treats explainability as a "
     "first-class deliverable alongside accuracy.")

section_h(doc, "2.4 Research Gap")
para(doc,
     "Much of the academic churn literature stops at comparative model "
     "accuracy. Fewer studies carry the model through to a deployed, "
     "explainable decision-support artifact with explicit threshold "
     "governance for business use. This paper addresses that gap by "
     "delivering the full pipeline — data validation, feature engineering, "
     "benchmarked and tuned models, three-way explainability, and an "
     "interactive dashboard — as a single, reproducible system.")

# ================================================== CHAPTER 3
chapter(doc, "CHAPTER 3: RESEARCH METHODOLOGY")

section_h(doc, "3.1 Research Design")
para(doc,
     "The study follows a quantitative, applied predictive-analytics design "
     "structured as: data ingestion and quality validation; exploratory data "
     "analysis; feature engineering; model benchmarking; hyperparameter "
     "tuning of the leading candidates; final evaluation on a held-out test "
     "set; and explainability analysis. All computation was performed in "
     "Python using scikit-learn (Pedregosa et al., 2011) and XGBoost, with "
     "every preprocessing step embedded in a serialized Pipeline so that "
     "training-time transformations are reproduced exactly at inference "
     "time.")

section_h(doc, "3.2 Data Description")
para(doc,
     "The dataset comprises 10,000 customer records with fourteen raw "
     "attributes from a European retail bank (snapshot year 2025). Data "
     "quality validation found zero missing values, zero duplicate records, "
     "and consistent binary encodings across all indicator fields.")
add_table(doc,
          ["Variable", "Type", "Description"],
          [
              ["CustomerId, Surname", "Identifier", "Dropped from modeling; excluded from all outputs"],
              ["CreditScore", "Numeric", "Creditworthiness, range 350–850"],
              ["Geography", "Categorical", "France (50.1%), Germany (25.1%), Spain (24.8%)"],
              ["Gender", "Categorical", "Male (54.6%), Female (45.4%)"],
              ["Age", "Numeric", "18–92 years, median 37"],
              ["Tenure", "Numeric", "Years with the bank, 0–10"],
              ["Balance", "Numeric", "Account balance; 36.2% of customers at zero"],
              ["NumOfProducts", "Numeric", "Products held, 1–4"],
              ["HasCrCard / IsActiveMember", "Binary", "Card ownership; activity indicator"],
              ["EstimatedSalary", "Numeric", "Annual salary proxy"],
              ["Exited", "Binary (target)", "1 = churned (20.37% prevalence)"],
          ],
          caption="Dataset variables and descriptive summary", table_no=1)

section_h(doc, "3.3 Feature Engineering")
para(doc,
     "Five derived features were constructed to encode behavioral signals "
     "not visible in the raw schema:")
add_table(doc,
          ["Feature", "Definition", "Behavioral hypothesis"],
          [
              ["BalanceSalaryRatio", "Balance / (Salary + 1)", "Wealth concentration relative to income"],
              ["ProductDensity", "Products / (Tenure + 1)", "Speed of product adoption"],
              ["EngagementScore", "IsActiveMember × Products", "Active engagement weighted by breadth"],
              ["AgeTenureRatio", "Age / (Tenure + 1)", "Life stage relative to relationship age"],
              ["ZeroBalance", "1 if Balance = 0", "Flags the dormant-balance segment (36.2%)"],
          ],
          caption="Engineered behavioral features", table_no=2)

section_h(doc, "3.4 Preprocessing Pipeline")
para(doc,
     "Identifiers and the constant Year column were dropped. Geography and "
     "Gender were one-hot encoded; all numeric features were standardized. "
     "Both transformations were wrapped with the classifier in a single "
     "scikit-learn Pipeline, persisted as one artifact, eliminating "
     "training–serving skew.")

section_h(doc, "3.5 Model Portfolio and Validation Strategy")
add_table(doc,
          ["Model", "Family", "Role in the benchmark"],
          [
              ["Logistic Regression", "Linear", "Interpretability baseline (class-weight balanced)"],
              ["Decision Tree", "Tree", "Single-tree reference (depth 8, balanced)"],
              ["Random Forest", "Bagged ensemble", "Variance reduction (300 trees, balanced)"],
              ["Gradient Boosting", "Boosted ensemble", "Sequential additive refinement"],
              ["XGBoost", "Boosted ensemble", "Regularized, histogram-based boosting"],
          ],
          caption="Candidate model portfolio", table_no=3)
para(doc,
     "Data were split 80/20 with stratification on the target, preserving "
     "the 20.37% churn prevalence in both partitions (n = 8,000 train; "
     "n = 2,000 test). The two leading models by test ROC-AUC were tuned "
     "with RandomizedSearchCV (15 candidate configurations each) under "
     "five-fold stratified cross-validation on the training partition only; "
     "the test set was touched exactly once per final model.")

section_h(doc, "3.6 Evaluation Metrics")
para(doc,
     "Six complementary metrics are reported: Accuracy, Precision, Recall, "
     "F1-score, ROC-AUC (threshold-free discrimination), and PR-AUC "
     "(average precision, more informative than ROC-AUC under the 20% class "
     "imbalance). ROC-AUC on the held-out test set is the selection "
     "criterion, with PR-AUC as tie-breaker.")

# ================================================== CHAPTER 4
chapter(doc, "CHAPTER 4: EXPLORATORY DATA ANALYSIS")

section_h(doc, "4.1 Target Distribution and Category-Level Churn")
add_figure(doc, "01_target_balance.png",
           "Target class balance — 2,037 churners (20.4%) among 10,000 customers",
           1, width=3.8)
add_figure(doc, "02_churn_by_category.png",
           "Churn rate by Geography, Gender, and activity status", 2)
para(doc,
     "Three categorical effects stand out. Germany churns at 32.4% — "
     "double France (16.1%) and Spain (16.7%) — despite hosting only a "
     "quarter of the base. Female customers churn at 25.1% versus 16.5% for "
     "male customers. Inactive members churn at 26.9% versus 14.3% for "
     "active members — the single most actionable behavioral gap in the "
     "portfolio.")

section_h(doc, "4.2 Distributional Differences by Outcome")
add_figure(doc, "03_distributions_by_target.png",
           "Feature distributions for retained versus churned customers", 3)
para(doc,
     "Age separates the classes most sharply: churn probability climbs from "
     "7.5% among customers aged 18–30 to 56.2% in the 51–60 band. "
     "Churned customers also carry systematically higher balances — an "
     "early indication that balance is not a loyalty signal in this "
     "portfolio.")

section_h(doc, "4.3 Correlation Structure and Interaction Effects")
add_figure(doc, "04_correlation_heatmap.png",
           "Correlation matrix across raw and engineered features", 4)
add_figure(doc, "05_age_products_heatmap.png",
           "Churn rate by age band and number of products", 5, width=5.0)
para(doc,
     "The age-by-products heatmap exposes the portfolio’s most striking "
     "interaction: churn among customers holding three or more products "
     "approaches certainty in every age band (82.7% at three products; 100% "
     "at four), while two-product customers are the safest segment at 7.6%. "
     "The conventional assumption that deeper product relationships imply "
     "stickier customers inverts beyond two products — the “three-product "
     "trap.”")
add_table(doc,
          ["Segment", "n", "Churn rate"],
          [
              ["Geography = Germany", "2,509", "32.4%"],
              ["Geography = France / Spain", "7,491", "16.1% / 16.7%"],
              ["Inactive members", "4,849", "26.9%"],
              ["Active members", "5,151", "14.3%"],
              ["NumOfProducts = 2", "4,590", "7.6%"],
              ["NumOfProducts ≥ 3", "326", "84.5% (combined)"],
              ["Age 51–60", "797", "56.2%"],
              ["Balance > 0", "6,383", "24.1%"],
              ["Balance = 0", "3,617", "13.8%"],
          ],
          caption="Churn rate by key customer segment", table_no=4)

# ================================================== CHAPTER 5
chapter(doc, "CHAPTER 5: MODEL DEVELOPMENT AND RESULTS")

section_h(doc, "5.1 Benchmark and Tuned Performance")
add_table(doc,
          ["Model", "Accuracy", "Precision", "Recall", "F1", "ROC-AUC", "PR-AUC"],
          [
              ["Gradient Boosting (tuned)", "0.870", "0.794", "0.484", "0.602", "0.870", "0.720"],
              ["Gradient Boosting", "0.869", "0.788", "0.484", "0.600", "0.869", "0.719"],
              ["XGBoost (tuned)", "0.871", "0.806", "0.479", "0.601", "0.868", "0.717"],
              ["XGBoost", "0.866", "0.772", "0.482", "0.593", "0.855", "0.700"],
              ["Random Forest", "0.862", "0.769", "0.457", "0.573", "0.853", "0.682"],
              ["Decision Tree", "0.760", "0.446", "0.742", "0.557", "0.804", "0.624"],
              ["Logistic Regression", "0.714", "0.388", "0.710", "0.502", "0.776", "0.469"],
          ],
          caption="Comparative model performance on the held-out test set (n = 2,000)",
          table_no=5, font_size=10)
add_figure(doc, "06_roc_curves.png",
           "ROC and precision–recall curves for all candidate models", 6)
para(doc,
     "Ensemble boosting dominates: both boosted families exceed 0.85 "
     "ROC-AUC, while the interpretable baselines trail by 7–10 points. The "
     "tuned Gradient Boosting classifier (learning rate 0.05, 200 "
     "estimators, maximum depth 3, subsample 1.0) attains the best test "
     "ROC-AUC (0.870) and PR-AUC (0.720). The tuned XGBoost model is "
     "statistically indistinguishable; Gradient Boosting was preferred for "
     "its smaller dependency footprint in deployment.")

section_h(doc, "5.2 Champion Model Diagnosis")
add_figure(doc, "07_confusion_matrix_best.png",
           "Confusion matrix of the tuned Gradient Boosting model at the 0.5 threshold",
           7, width=4.0)
para(doc,
     "At the default 0.5 threshold the champion model favors precision "
     "(0.794) over recall (0.484): four of five customers it flags are "
     "genuine churn risks, minimizing wasted retention spend, at the cost "
     "of catching roughly half of eventual churners. This trade-off is a "
     "policy choice, not a model property: lowering the threshold to 0.30 "
     "— the dashboard’s medium-risk boundary — captures materially more "
     "churners at higher outreach cost. Threshold governance is assigned to "
     "the Retention Committee (Recommendation R6).")

# ================================================== CHAPTER 6
chapter(doc, "CHAPTER 6: MODEL EXPLAINABILITY")

section_h(doc, "6.1 Global Feature Importance")
add_figure(doc, "09_feature_importance.png",
           "Gain-based feature importance of the champion model (top 15)", 8)
add_table(doc,
          ["Rank", "Feature", "Importance (gain share)"],
          [
              ["1", "Age", "0.390"],
              ["2", "NumOfProducts", "0.285"],
              ["3", "EngagementScore (engineered)", "0.084"],
              ["4", "Balance", "0.065"],
              ["5", "Geography_Germany", "0.050"],
              ["6", "IsActiveMember", "0.043"],
              ["7", "BalanceSalaryRatio (engineered)", "0.023"],
              ["8", "CreditScore", "0.018"],
          ],
          caption="Top churn drivers by gain importance", table_no=6)

section_h(doc, "6.2 SHAP Analysis")
add_figure(doc, "10_shap_summary.png",
           "SHAP beeswarm summary — direction and magnitude of each feature’s impact",
           9, width=5.2)
add_figure(doc, "11_shap_bar.png",
           "Mean absolute SHAP value per feature", 10, width=5.0)
para(doc,
     "SHAP corroborates the gain ranking and adds direction: high age and "
     "three-plus product holdings push predictions toward churn; active "
     "membership pulls them away; German residency raises risk while France "
     "and Spain are approximately neutral. The engineered EngagementScore "
     "ranks third overall, validating the feature-engineering hypothesis "
     "that activity interacted with product breadth carries signal beyond "
     "either raw variable.")

section_h(doc, "6.3 Partial Dependence")
add_figure(doc, "12_partial_dependence.png",
           "Partial-dependence profiles for the five dominant drivers", 11)
para(doc,
     "Partial dependence confirms the non-linear age curve (risk "
     "accelerating between 40 and 60), the two-product safety trough, and "
     "the protective effect of activity — consistent across all three "
     "explainability lenses.")

# ================================================== CHAPTER 7
chapter(doc, "CHAPTER 7: FINDINGS AND RECOMMENDATIONS")

section_h(doc, "7.1 Consolidated Findings")
bullets(doc, [
    "F1 — Age is the dominant churn driver; the 51–60 band churns at 56.2%.",
    "F2 — Product depth is non-monotonic: two products is the retention sweet spot (7.6%); three or more signals near-certain churn (83–100%).",
    "F3 — Engagement is the cheapest lever: inactive members churn at 26.9% vs 14.3% for active members.",
    "F4 — Germany churns at double the rate of France and Spain; the effect is not explained by observable demographics or balances.",
    "F5 — The tuned Gradient Boosting model provides 0.870 ROC-AUC with fully explainable predictions.",
])

section_h(doc, "7.2 Managerial Recommendations")
add_table(doc,
          ["#", "Recommendation", "Grounding"],
          [
              ["R1", "Cap automated cross-sell at two products; require a satisfaction check before any third product", "F2"],
              ["R2", "Launch an activation campaign (nudges, relationship-manager outreach) on the inactive cohort — the highest expected ROI lever", "F3"],
              ["R3", "Commission a Germany-specific diagnostic covering pricing, service quality, and product fit", "F4"],
              ["R4", "Design a pre-retirement wealth-transition proposition for customers aged 50–60", "F1"],
              ["R5", "Operationalize monthly risk scoring; route the top decile to relationship managers; measure uplift with a holdout control group", "F5"],
              ["R6", "Assign decision rights over the alert threshold to the Retention Committee; the 0.30 boundary doubles churner capture at manageable cost", "§5.2"],
          ],
          caption="Managerial recommendations mapped to findings", table_no=7,
          font_size=10)

# ================================================== CHAPTER 8
chapter(doc, "CHAPTER 8: LIMITATIONS AND FUTURE SCOPE")
bullets(doc, [
    "Cross-sectional snapshot: the data cannot support survival analysis (when a customer will churn) or capture time-varying macro effects; a panel extension is the natural next step.",
    "Recall at the default threshold is 0.484; a cost-sensitive optimization aligned with the bank’s actual retention economics should precede full-scale rollout.",
    "Fairness has not been audited: Geography (and correlated attributes) is a strong driver, and an adverse-impact assessment on protected characteristics is required before deployment in a supervised jurisdiction.",
    "Probabilities are uncalibrated; Platt or isotonic calibration is advised where scores feed budget allocation.",
    "The three-plus-product churn pattern is an association in snapshot data; product-add timestamps are needed to separate distress cross-selling from causal product effects.",
])

# ================================================== CHAPTER 9
chapter(doc, "CHAPTER 9: CONCLUSION")
para(doc,
     "This research delivers a complete, deployable churn-intelligence "
     "system for a European retail bank. A disciplined benchmark of five "
     "algorithms, tuned under stratified cross-validation, produced a "
     "champion Gradient Boosting model with 0.870 test ROC-AUC whose every "
     "prediction is decomposable into feature contributions. The "
     "explainability analysis converged on four drivers — age, product "
     "depth, engagement, and German residency — and surfaced the "
     "portfolio’s most consequential behavioral insight: cross-selling "
     "beyond two products is associated with near-certain attrition. "
     "Combined with the interactive dashboard and the recommended threshold "
     "governance, the system equips the bank to shift retention practice "
     "from reactive reporting to targeted, explainable, and measurable "
     "intervention.")

# ================================================== REFERENCES
chapter(doc, "REFERENCES")
for ref in [
    "Breiman, L. (2001). Random forests. Machine Learning, 45(1), 5–32.",
    "Chen, T., & Guestrin, C. (2016). XGBoost: A scalable tree boosting "
    "system. In Proceedings of the 22nd ACM SIGKDD International Conference "
    "on Knowledge Discovery and Data Mining (pp. 785–794). ACM.",
    "Friedman, J. H. (2001). Greedy function approximation: A gradient "
    "boosting machine. The Annals of Statistics, 29(5), 1189–1232.",
    "Gupta, S., Lehmann, D. R., & Stuart, J. A. (2004). Valuing customers. "
    "Journal of Marketing Research, 41(1), 7–18.",
    "Lemmens, A., & Croux, C. (2006). Bagging and boosting classification "
    "trees to predict churn. Journal of Marketing Research, 43(2), 276–286.",
    "Lundberg, S. M., & Lee, S.-I. (2017). A unified approach to "
    "interpreting model predictions. In Advances in Neural Information "
    "Processing Systems 30 (pp. 4765–4774).",
    "Neslin, S. A., Gupta, S., Kamakura, W., Lu, J., & Mason, C. H. (2006). "
    "Defection detection: Measuring and understanding the predictive "
    "accuracy of customer churn models. Journal of Marketing Research, "
    "43(2), 204–211.",
    "Pedregosa, F., Varoquaux, G., Gramfort, A., Michel, V., Thirion, B., "
    "Grisel, O., … Duchesnay, É. (2011). Scikit-learn: Machine learning "
    "in Python. Journal of Machine Learning Research, 12, 2825–2830.",
    "Reichheld, F. F., & Sasser, W. E., Jr. (1990). Zero defections: "
    "Quality comes to services. Harvard Business Review, 68(5), 105–111.",
    "Verbeke, W., Dejaeger, K., Martens, D., Hur, J., & Baesens, B. (2012). "
    "New insights into churn prediction in the telecommunication sector: A "
    "profit driven data mining approach. European Journal of Operational "
    "Research, 218(1), 211–229.",
]:
    reference(doc, ref)

# ================================================== APPENDICES
chapter(doc, "APPENDIX A: THE DEPLOYED DASHBOARD")
para(doc,
     "The model is operationalized as a five-module Streamlit web "
     "application deployed on Streamlit Community Cloud (repository: "
     "github.com/shubhangi-ppandey/bank-churn-prediction).")
bullets(doc, [
    "Risk Calculator — enter a customer profile; receive churn probability, risk band (Low < 30% ≤ Medium < 60% ≤ High), and a recommended action.",
    "Probability Distribution — scores the full 10,000-customer portfolio; risk-band counts and top-25 high-risk list.",
    "Feature Importance — interactive gain chart with SHAP and partial-dependence visuals.",
    "What-If Simulator — perturb a real customer’s engagement, products, or balance and observe the marginal probability change.",
    "Model Performance — full metrics table, ROC/PR curves, confusion matrix, and EDA gallery.",
])

chapter(doc, "APPENDIX B: TECHNOLOGY STACK AND REPRODUCIBILITY")
add_table(doc,
          ["Layer", "Technology"],
          [
              ["Language / runtime", "Python 3.12 (pinned in runtime.txt)"],
              ["Data handling", "pandas, NumPy"],
              ["Modeling", "scikit-learn (Pipeline, GradientBoostingClassifier), XGBoost"],
              ["Explainability", "SHAP (TreeExplainer), sklearn partial dependence"],
              ["Visualization", "Matplotlib, Seaborn, Plotly"],
              ["Application / deployment", "Streamlit; Streamlit Community Cloud; pinned requirements.txt"],
              ["Version control", "Git / GitHub (github.com/shubhangi-ppandey/bank-churn-prediction)"],
          ],
          caption="Technology stack", table_no=8)
para(doc,
     "All figures and tables in this paper are regenerated end-to-end by "
     "three scripts (src/eda.py, src/train.py, src/explain.py) against the "
     "committed dataset, guaranteeing full reproducibility of every "
     "reported number.")

doc.save(OUT)
print(f"Saved: {OUT}")
